import os
from datetime import datetime, timedelta

import pandas as pd
from docx import Document

from src.drive_utils import get_drive_service, upload_file_to_subfolder

ROOT_DRIVE_FOLDER_ID = os.environ.get("GOOGLE_DRIVE_FOLDER_ID")
OUTPUT_DIR = "output"


def list_recent_csv_files(service, root_folder_id, days=8):
    # This lists CSV files under all accessible Drive folders created by this automation.
    # It searches by file name pattern because Drive API cannot easily recursive-list in one query.
    query = "name contains 'PEARL_commonality_news_' and name contains '.csv' and trashed=false"
    res = service.files().list(q=query, fields="files(id, name, createdTime)", pageSize=1000).execute()
    files = res.get("files", [])
    cutoff = datetime.utcnow() - timedelta(days=days)
    recent = []
    for f in files:
        try:
            created = datetime.strptime(f["createdTime"].replace("Z", ""), "%Y-%m-%dT%H:%M:%S.%f")
        except ValueError:
            created = datetime.strptime(f["createdTime"].replace("Z", ""), "%Y-%m-%dT%H:%M:%S")
        if created >= cutoff:
            recent.append(f)
    return recent


def download_file(service, file_id, local_path):
    from googleapiclient.http import MediaIoBaseDownload
    import io

    request = service.files().get_media(fileId=file_id)
    fh = io.FileIO(local_path, "wb")
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()


def main():
    if not ROOT_DRIVE_FOLDER_ID:
        raise RuntimeError("Missing GOOGLE_DRIVE_FOLDER_ID secret/environment variable.")

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    service = get_drive_service()
    recent_files = list_recent_csv_files(service, ROOT_DRIVE_FOLDER_ID, days=8)

    frames = []
    for f in recent_files:
        local_csv = os.path.join(OUTPUT_DIR, f["name"])
        download_file(service, f["id"], local_csv)
        frames.append(pd.read_csv(local_csv))

    today = datetime.utcnow().strftime("%Y-%m-%d")
    report_path = os.path.join(OUTPUT_DIR, f"PEARL_Weekly_Commonality_Report_{today}.docx")

    doc = Document()
    doc.add_heading("PEARL Weekly Commonality Market and News Report", 0)
    doc.add_paragraph(f"Generated on: {today}")

    if not frames:
        doc.add_paragraph("No daily CSV files found for this week.")
        doc.save(report_path)
        upload_file_to_subfolder(report_path, ROOT_DRIVE_FOLDER_ID, "03_Weekly_Report")
        return

    data = pd.concat(frames, ignore_index=True)
    data = data.drop_duplicates(subset=["url"])
    data = data.sort_values("relevance_score", ascending=False)

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This weekly report summarizes Cambodia news and global trends related to PEARL commonality crops: "
        "mango, cashew, rice and vegetables. It covers market information, prices, exports, production, "
        "climate risks, policy, investment, processing, GHG and sustainability."
    )

    doc.add_heading("2. Summary Statistics", level=1)
    doc.add_paragraph(f"Total unique references collected this week: {len(data)}")

    doc.add_heading("3. Key Findings by Crop", level=1)
    for crop in ["Mango", "Cashew", "Rice", "Vegetables"]:
        doc.add_heading(crop, level=2)
        crop_df = data[data["crop"] == crop].head(10)
        if crop_df.empty:
            doc.add_paragraph("No relevant articles collected this week.")
        else:
            for _, row in crop_df.iterrows():
                doc.add_paragraph(
                    f"- {row.get('title','')} | Source: {row.get('source','')} | "
                    f"Group: {row.get('source_group','')} | Topic: {row.get('topic','')} | "
                    f"Relevance: {row.get('pearl_relevance','')}"
                )

    doc.add_heading("4. Strategic Implications for PEARL", level=1)
    doc.add_paragraph(
        "Use these references to support PEARL monitoring of crop value chains, market access, climate risk, "
        "investment targeting, farmer organization support, agroecological monitoring and policy tracking."
    )

    doc.add_heading("5. References and Citations", level=1)
    for i, (_, row) in enumerate(data.iterrows(), start=1):
        doc.add_paragraph(f"[{i}] {row.get('citation','')}")

    doc.save(report_path)
    upload_file_to_subfolder(report_path, ROOT_DRIVE_FOLDER_ID, "03_Weekly_Report")
    print(f"Weekly report complete: {report_path}")


if __name__ == "__main__":
    main()
